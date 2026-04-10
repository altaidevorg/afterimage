"""
File operation utilities for the demo UI.
"""

import os
import shutil
import zipfile
import pandas as pd
import json
from pathlib import Path
from typing import Optional, List, Union, Any
from html.parser import HTMLParser

from afterimage import InMemoryDocumentProvider
from afterimage.providers import JSONLDocumentProvider


# --- Text Processing Classes ---


class HTMLTextExtractor(HTMLParser):
    """Extract plain text from HTML content."""

    def __init__(self):
        super().__init__()
        self.result = []

    def handle_data(self, data):
        if data.strip():
            self.result.append(data.strip())

    def get_text(self):
        return "\n".join(self.result)


class SimpleTextSplitter:
    """Simple text chunking with overlap."""

    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def split_text(self, text: str) -> List[str]:
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)

            # Move start pointer
            start += self.chunk_size - self.overlap

            # Avoid infinite loop if overlap >= chunk_size
            if self.overlap >= self.chunk_size:
                start = end

        return chunks


# --- File Operations ---


def merge_dataset_files(
    source_files: List[Union[str, Any]],
    dest_dir: str,
    filename: str = "toolcalldataset.jsonl",
) -> str:
    """
    Merge multiple dataset files into a single temporary training file.

    Args:
        source_files: List of file paths or file objects
        dest_dir: Destination directory path
        filename: Target filename (default: toolcalldataset.jsonl)

    Returns:
        Path to the merged file
    """
    os.makedirs(dest_dir, exist_ok=True)
    dest_path = os.path.join(dest_dir, filename)

    with open(dest_path, "wb") as outfile:
        for f in source_files:
            source_path = f if isinstance(f, str) else f.name
            with open(source_path, "rb") as infile:
                shutil.copyfileobj(infile, outfile)
                # Ensure newline between files if missing
                outfile.write(b"\n")

    return dest_path


def copy_dataset_file(
    source_file, dest_dir: str, filename: str = "toolcalldataset.jsonl"
) -> str:
    """
    Copy a dataset file to the training data directory.

    Args:
        source_file: File path string or file object with .name attribute
        dest_dir: Destination directory path
        filename: Target filename (default: toolcalldataset.jsonl)

    Returns:
        Path to the copied file
    """
    os.makedirs(dest_dir, exist_ok=True)
    dest_path = os.path.join(dest_dir, filename)

    # Handle both string paths and file objects
    source_path = source_file if isinstance(source_file, str) else source_file.name
    shutil.copy(source_path, dest_path)
    return dest_path


def create_model_zip(model_dir: str, output_name: str = "model.zip") -> Optional[str]:
    """
    Create a zip file from a model directory.

    Args:
        model_dir: Path to the model directory
        output_name: Name for the output zip file

    Returns:
        Path to the created zip file, or None if failed
    """
    if not os.path.exists(model_dir):
        return None

    if not os.path.isdir(model_dir):
        return None

    try:
        zip_path = os.path.join(os.path.dirname(model_dir), output_name)

        # Remove old zip if exists
        if os.path.exists(zip_path):
            os.remove(zip_path)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(model_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, os.path.dirname(model_dir))
                    zipf.write(file_path, arcname)

        return zip_path
    except Exception as e:
        print(f"Error creating zip: {e}")
        return None


def clean_temp_files(*file_paths):
    """
    Remove temporary files.

    Args:
        *file_paths: Variable number of file paths to remove
    """
    for path in file_paths:
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except Exception as e:
                print(f"Warning: Could not remove {path}: {e}")


# --- Document Processing ---


def extract_text_from_file(file_path: Path) -> str:
    """
    Extract text content from various file formats.

    Supports: .docx, .rtf, .html, .htm, and plain text files.

    Args:
        file_path: Path to the file

    Returns:
        Extracted text content

    Raises:
        ImportError: If required library is not installed
        ValueError: If no text could be extracted
    """
    ext = file_path.suffix.lower()
    text = ""

    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for .docx files. Please install it."
            )

        doc = docx.Document(file_path)
        # Extract paragraphs
        paras = [para.text for para in doc.paragraphs]
        # Extract tables
        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                tables.append(row_text)

        text = "\n".join(paras + tables)

    elif ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf is required for .rtf files. Please install it.")

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        text = rtf_to_text(content)

    elif ext == ".html" or ext == ".htm":
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        parser = HTMLTextExtractor()
        parser.feed(content)
        text = parser.get_text()

    else:
        # Fallback for plain text-like files
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

    if not text.strip():
        raise ValueError(f"No text could be extracted from {file_path.name}")

    return text


def create_document_provider_from_file(
    file_path: str, key: str = "text", chunk_size: int = 1000, chunk_overlap: int = 100
):
    """
    Creates an appropriate DocumentProvider based on the file extension.

    Supported formats:
    - Structured:
        - .csv: Uses pandas to read the specified column
        - .tsv: Uses pandas with tab separator
        - .jsonl: Uses JSONLDocumentProvider with the specified key
        - .txt: Manual lines (split by newline) - Treated as structured/pre-chunked

    - Raw Documents (will be chunked):
        - .docx: Word documents
        - .rtf: Rich Text Format
        - .html: HTML files

    Args:
        file_path: Path to the file
        key: Column/key name for structured formats
        chunk_size: Size of text chunks for raw documents
        chunk_overlap: Overlap between chunks

    Returns:
        DocumentProvider instance

    Raises:
        ValueError: If file format is unsupported or column not found
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    # Structured / Pre-chunked formats
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            lines = [line.strip() for line in f if line.strip()]
        return InMemoryDocumentProvider(lines)

    elif ext == ".csv":
        df = pd.read_csv(file_path)
        if key not in df.columns:
            raise ValueError(
                f"Column '{key}' not found in CSV. Available columns: {list(df.columns)}"
            )
        texts = df[key].dropna().astype(str).tolist()
        return InMemoryDocumentProvider(texts)

    elif ext == ".tsv":
        df = pd.read_csv(file_path, sep="\t")
        if key not in df.columns:
            raise ValueError(
                f"Column '{key}' not found in TSV. Available columns: {list(df.columns)}"
            )
        texts = df[key].dropna().astype(str).tolist()
        return InMemoryDocumentProvider(texts)

    elif ext == ".jsonl":
        return JSONLDocumentProvider(path_pattern=file_path, content_key=key)

    # Raw formats requiring extraction and chunking
    elif ext in [".docx", ".rtf", ".html", ".htm"]:
        raw_text = extract_text_from_file(path)
        splitter = SimpleTextSplitter(chunk_size=chunk_size, overlap=chunk_overlap)
        chunks = splitter.split_text(raw_text)
        # Filter empty chunks
        chunks = [c.strip() for c in chunks if c.strip()]
        return InMemoryDocumentProvider(chunks)

    else:
        raise ValueError(f"Unsupported file extension: {ext}")
