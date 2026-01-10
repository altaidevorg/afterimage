from html.parser import HTMLParser
import pandas as pd
from typing import List
from pathlib import Path
from afterimage import InMemoryDocumentProvider
from afterimage.providers import JSONLDocumentProvider


class HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.result = []

    def handle_data(self, data):
        if data.strip():
            self.result.append(data.strip())

    def get_text(self):
        return "\n".join(self.result)


class SimpleTextSplitter:
    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def split_text(self, text: str) -> List[str]:
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)

            # Move start pointer
            start += self.chunk_size - self.overlap

            # Avoid infinite loop if overlap >= chunk_size (shouldn't happen with sane defaults)
            if self.overlap >= self.chunk_size:
                start = end

        return chunks


def extract_text_from_file(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    text = ""

    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for .docx files. Please install it."
            )

        doc = docx.Document(file_path)
        # Extract paragraphs
        paras = [para.text for para in doc.paragraphs]
        # Extract tables
        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                tables.append(row_text)

        text = "\n".join(paras + tables)

    elif ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf is required for .rtf files. Please install it.")

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        text = rtf_to_text(content)

    elif ext == ".html" or ext == ".htm":
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        parser = HTMLTextExtractor()
        parser.feed(content)
        text = parser.get_text()

    else:
        # Fallback for plain text-like files
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

    if not text.strip():
        raise ValueError(f"No text could be extracted from {file_path.name}")

    return text


def create_document_provider_from_file(
    file_path: str, key: str = "text", chunk_size: int = 1000, chunk_overlap: int = 100
):
    """
    Creates an appropriate DocumentProvider based on the file extension.

    Supported formats:
    - Structured:
        - .csv: Uses pandas to read the specified column
        - .tsv: Uses pandas with tab separator
        - .jsonl: Uses JSONLDocumentProvider with the specified key
        - .txt: Manual lines (split by newline) - Treated as structured/pre-chunked

    - Raw Documents (will be chunked):
        - .docx: Word documents
        - .rtf: Rich Text Format
        - .html: HTML files
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    # Structured / Pre-chunked formats
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            lines = [line.strip() for line in f if line.strip()]
        return InMemoryDocumentProvider(lines)

    elif ext == ".csv":
        df = pd.read_csv(file_path)
        if key not in df.columns:
            raise ValueError(
                f"Column '{key}' not found in CSV. Available columns: {list(df.columns)}"
            )
        texts = df[key].dropna().astype(str).tolist()
        return InMemoryDocumentProvider(texts)

    elif ext == ".tsv":
        df = pd.read_csv(file_path, sep="\t")
        if key not in df.columns:
            raise ValueError(
                f"Column '{key}' not found in TSV. Available columns: {list(df.columns)}"
            )
        texts = df[key].dropna().astype(str).tolist()
        return InMemoryDocumentProvider(texts)

    elif ext == ".jsonl":
        return JSONLDocumentProvider(path_pattern=file_path, content_key=key)

    # Raw formats requiring extraction and chunking
    elif ext in [".docx", ".rtf", ".html", ".htm"]:
        raw_text = extract_text_from_file(path)
        splitter = SimpleTextSplitter(chunk_size=chunk_size, overlap=chunk_overlap)
        chunks = splitter.split_text(raw_text)
        # Filter empty chunks
        chunks = [c.strip() for c in chunks if c.strip()]
        return InMemoryDocumentProvider(chunks)

    else:
        raise ValueError(f"Unsupported file extension: {ext}")
