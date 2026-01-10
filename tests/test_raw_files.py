import os
import unittest
from pathlib import Path

try:
    import docx
except ImportError:
    docx = None

from examples.demo_ui.utils import create_document_provider_from_file


class TestRawFileParsing(unittest.TestCase):
    def setUp(self):
        self.files_to_remove = []

    def tearDown(self):
        for f in self.files_to_remove:
            if os.path.exists(f):
                os.remove(f)

    def test_docx(self):
        if docx is None:
            print("Skipping DOCX test (python-docx not installed)")
            return

        filename = "test_doc.docx"
        doc = docx.Document()
        doc.add_paragraph("Hello world.")
        doc.add_paragraph("This is a test document.")
        # Add enough text to trigger chunking if chunk_size is small,
        # but here we rely on default 1000.
        # Let's just check extraction first.
        doc.save(filename)
        self.files_to_remove.append(filename)

        provider = create_document_provider_from_file(filename)
        docs = provider.get_all()
        # Should be one chunk if small enough
        text = " ".join([d.text for d in docs])
        self.assertIn("Hello world.", text)
        self.assertIn("This is a test document.", text)

    def test_html(self):
        filename = "test_doc.html"
        content = "<html><body><h1>Title</h1><p>Paragraph 1</p><p>Paragraph 2</p></body></html>"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(content)
        self.files_to_remove.append(filename)

        provider = create_document_provider_from_file(filename)
        docs = provider.get_all()
        text = " ".join([d.text for d in docs])
        # HTMLTextExtractor joins with newlines
        self.assertIn("Title", text)
        self.assertIn("Paragraph 1", text)
        self.assertIn("Paragraph 2", text)

    def test_rtf(self):
        # minimal RTF
        content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Courier;}} \f0\fs20 Hello World! This is RTF.}"
        filename = "test_doc.rtf"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(content)
        self.files_to_remove.append(filename)

        # This might fail if striprtf is not installed or if my minimal RTF is too minimal
        try:
            provider = create_document_provider_from_file(filename)
            docs = provider.get_all()
            text = " ".join([d.text for d in docs])
            self.assertIn("Hello World!", text)
            self.assertIn("This is RTF.", text)
        except ImportError:
            print("Skipping RTF test (striprtf not installed)")


if __name__ == "__main__":
    unittest.main()
